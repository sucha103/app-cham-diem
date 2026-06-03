import os
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_guide_docx():
    doc = docx.Document()
    
    # Thiết lập lề trang
    sections = doc.sections
    for section in sections:
        section.top_margin = docx.shared.Inches(1)
        section.bottom_margin = docx.shared.Inches(1)
        section.left_margin = docx.shared.Inches(1)
        section.right_margin = docx.shared.Inches(1)

    # Hàm phụ trợ thiết lập định dạng chữ
    def set_font(run, name="Segoe UI", size=11, bold=False, italic=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    # Hàm thêm tiêu đề chính
    def add_main_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(20)
        run = p.add_run(text)
        set_font(run, size=18, bold=True, color=RGBColor(44, 62, 80))
        return p

    # Hàm thêm tiêu đề cấp 1
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_font(run, size=14, bold=True, color=RGBColor(41, 128, 185))
        return p

    # Hàm thêm tiêu đề cấp 2
    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_font(run, size=12, bold=True, color=RGBColor(52, 73, 94))
        return p

    # Hàm thêm đoạn văn thường
    def add_body_text(text, bold_prefix="", space_after=6, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            set_font(run_prefix, bold=True)
            
        run = p.add_run(text)
        set_font(run, italic=italic)
        return p

    # Hàm thêm danh sách
    def add_bullet_point(text, bold_prefix="", space_after=3):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            set_font(run_prefix, bold=True)
            
        run = p.add_run(text)
        set_font(run)
        return p

    # Hàm thêm ô hộp ghi chú (Callout box)
    def add_callout_box(text, title="LƯU Ý QUAN TRỌNG"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        # Đặt chiều rộng cho bảng
        for row in tbl.rows:
            for cell in row.cells:
                cell.width = docx.shared.Inches(6.5)
                
                # Định dạng viền và màu nền bằng XML trực tiếp
                tcPr = cell._tc.get_or_add_tcPr()
                # Màu nền xám nhạt/xanh nhạt
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FA"/>')
                tcPr.append(shd)
                
                # Chỉ hiển thị viền bên trái dày màu xanh dương đậm
                borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="2980B9"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
                tcPr.append(borders)
                
                # Thêm padding
                mar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="120" w:type="dxa"/><w:bottom w:w="120" w:type="dxa"/><w:left w:w="200" w:type="dxa"/><w:right w:w="200" w:type="dxa"/></w:tcMar>')
                tcPr.append(mar)
                
                # Viết nội dung trong ô
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                
                run_title = p.add_run(f"🔔 {title}\n")
                set_font(run_title, bold=True, color=RGBColor(41, 128, 185))
                
                run_text = p.add_run(text)
                set_font(run_text, size=10, italic=True)
                
        doc.add_paragraph().paragraph_format.space_after = Pt(6) # Tạo khoảng trống sau bảng

    # --- BẮT ĐẦU VIẾT NỘI DUNG ---
    add_main_title("HƯỚNG DẪN SỬ DỤNG MINI APP CHẤM BÀI 28 NGÀY")
    
    add_body_text("Tài liệu này hướng dẫn chi tiết cách thiết lập, cấu hình và vận hành phần mềm hỗ trợ chấm bài viết tự động bằng AI kết hợp Google Sheets của cô Hạnh.", italic=True)
    
    # ----------------------------------------------------
    add_heading_1("I. Giới thiệu chung")
    add_body_text("Phần mềm chấm bài 28 ngày là một ứng dụng Web tĩnh (Static Web App) hoạt động ngay trên trình duyệt web cá nhân. Ứng dụng giúp giáo viên đọc nhanh các bài làm của học viên định dạng Word (.docx), phân tích sâu sắc các khía cạnh tư duy và hành văn bằng trí tuệ nhân tạo (Gemma AI), rồi tự động đồng bộ hóa kết quả lên Google Sheets để quản lý tập trung.")
    
    add_body_text("Các đặc điểm nổi bật:")
    add_bullet_point("Không cần cài đặt phần mềm phức tạp, mở link là dùng được ngay trên máy tính, iPad, điện thoại.", "Tiện lợi:")
    add_bullet_point("API Key và URL Sheets được lưu an toàn tại bộ nhớ cục bộ trình duyệt (localStorage), không gửi đi bên thứ ba.", "Bảo mật tuyệt đối:")
    add_bullet_point("Giao diện kính mờ sang trọng, hỗ trợ giao diện Sáng/Tối và co giãn tương thích hoàn hảo cho thiết bị di động.", "Trải nghiệm cao cấp:")
    
    # ----------------------------------------------------
    add_heading_1("II. Hướng dẫn cấu hình lần đầu")
    add_body_text("Chị chỉ cần cấu hình bước này duy nhất một lần đầu tiên khi mở trang web trên một thiết bị mới. Trình duyệt sẽ tự động ghi nhớ cho các lần sử dụng tiếp theo.")
    
    add_heading_2("Bước 1: Chuẩn bị các khóa cấu hình cần thiết")
    add_bullet_point("Truy cập vào trang web aistudio.google.com (đăng nhập bằng tài khoản Gmail của chị) để tạo một khóa API miễn phí.", "1. Google AI Studio API Key: ")
    add_bullet_point("Là đường dẫn URL Web App nhận được sau khi chị nhấn nút 'Triển khai' (Deploy) dự án Google Apps Script (từ file Code.gs) dưới quyền truy cập 'Mọi người' (Anyone).", "2. Google Apps Script Web App URL: ")
    
    add_heading_2("Bước 2: Điền cấu hình vào trang web")
    add_body_text("1. Mở trang web chấm bài trực tuyến tại địa chỉ: https://sucha103.github.io/app-cham-diem/ (Hoặc dùng link dự phòng: https://co-hanh-cham-diem.surge.sh).")
    add_body_text("2. Nhấp vào nút Cấu hình (biểu tượng hình bánh răng ⚙️) ở góc trên bên phải màn hình.")
    add_body_text("3. Điền khóa API và đường dẫn URL Web App vào hai ô nhập liệu tương ứng.")
    add_body_text("4. Nhấp nút Lưu cấu hình. Trang web sẽ thông báo lưu thành công và tự động đóng hộp thoại.")

    add_callout_box(
        "Nếu chị chủ động xoá dữ liệu lịch sử duyệt web (Clear Browsing Cache) hoặc chuyển sang dùng thiết bị khác (như điện thoại mới), chị mới cần vào lại nút bánh răng để dán lại 2 dòng cấu hình này.",
        "MẸO NHỎ VỀ BỘ NHỚ CẤU HÌNH"
    )

    # ----------------------------------------------------
    add_heading_1("III. Các bước chấm bài hàng ngày")
    add_body_text("Khi chấm bài cho học viên hàng ngày, chị chỉ cần thực hiện theo các bước cực kỳ trực quan dưới đây:")
    
    add_heading_2("Bước 1: Tải file bài làm lên ứng dụng")
    add_body_text("Chị chỉ cần kéo file Word (.docx) bài làm của học viên thả vào ô nét đứt (Dropzone) ở trên cùng, hoặc nhấp vào đó để chọn file từ máy tính.")
    add_body_text("Hệ thống sẽ tự động phân tích tên file hoặc nội dung bên trong file để tự điền các thông tin sau:")
    add_bullet_point("Học viên Nguyễn Thị Bình nộp file tên HV999_NguyenThiBinh_Ngay5.docx sẽ tự tách ra Mã số: HV999 và Họ tên: Nguyễn Thị Bình.", "Tự nhận dạng Họ tên & Mã học viên: ")
    add_bullet_point("Tự động chọn đúng Ngày học (từ ngày 1 đến ngày 28) và hiển thị bảng Câu hỏi định hướng tư duy phản tư của ngày đó bên dưới để chị đối chiếu.", "Tự nhận dạng Ngày học: ")
    add_bullet_point("Trích xuất và đổ toàn bộ văn bản thô từ file Word vào khung soạn thảo nội dung.", "Đọc nội dung văn bản: ")
    
    add_heading_2("Bước 2: Chọn kỹ năng AI muốn đánh giá (Gemma Skills)")
    add_body_text("Trên giao diện chính hiển thị 6 ô tích chọn (công tắc bật/tắt) tương ứng với 6 kỹ năng chấm bài chuyên sâu. Chị có thể tích chọn những kỹ năng chị muốn AI tập trung phân tích cho bài làm này:")
    add_bullet_point("Đánh giá chiều sâu tự nhận thức và độ trung thực với bản thân.", "1. Tư duy phản tư: ")
    add_bullet_point("Kiểm tra tính mạch lạc, cấu trúc lập luận và sự liên kết giữa các ý.", "2. Logic & Lập luận: ")
    add_bullet_point("Rà soát lỗi chính tả, câu văn dài dòng và đề xuất cách diễn đạt cô đọng.", "3. Hành văn & Ngôn từ: ")
    add_bullet_point("Đánh giá việc dùng chi tiết thực tế, hình ảnh trực quan thay vì kể chung chung.", "4. Hình ảnh hóa (Show, Don't Tell): ")
    add_bullet_point("Đánh giá dấu ấn cá nhân, tính chân thực và sự độc bản trong giọng văn.", "5. Cá tính & Giọng điệu (Voice): ")
    add_bullet_point("Đánh giá độ chạm, cảm xúc truyền tải và khả năng đồng cảm với người đọc.", "6. Độ chạm & Thấu cảm: ")

    add_heading_2("Bước 3: Nhấn Bắt đầu chấm bài")
    add_body_text("Sau khi chọn xong các kỹ năng cần thiết, chị nhấp nút Bắt đầu chấm bài. AI sẽ xử lý bài viết trong khoảng từ 5 - 10 giây và trả về bản nhận xét nháp đầy đủ cấu trúc: Tổng quan, Phân tích (Điểm tốt, Điểm cần cải thiện), và Lời khuyên cụ thể.")
    
    add_heading_2("Bước 4: Xem xét và tinh chỉnh nhận xét (AI Rewrite Loop)")
    add_body_text("Nếu bản nhận xét nháp của AI đã hoàn hảo, chị chỉ cần bấm Lưu & Cập nhật Sheet.")
    add_body_text("Nếu muốn điều chỉnh nhận xét, chị có hai lựa chọn cực kỳ linh hoạt:")
    add_bullet_point("Nhấp nút 'Sửa nhận xét' ngay dưới khung kết quả để tự gõ/chỉnh sửa văn bản theo ý mình trước khi lưu.", "Cách 1: Tự sửa tay: ")
    add_bullet_point("Nhấp nút 'Góp ý cho AI', gõ yêu cầu điều chỉnh (Ví dụ: 'Hãy viết thêm lời khích lệ ở cuối và cho em ấy một bài tập nhỏ viết nhật ký nhé') rồi bấm 'AI Viết lại'. AI sẽ viết lại bản nhận xét mới tối ưu hơn dựa trên bài làm cũ và góp ý của chị.", "Cách 2: Yêu cầu AI viết lại: ")

    add_heading_2("Bước 5: Lưu dữ liệu và kiểm tra trên Google Sheets")
    add_body_text("Sau khi bấm Lưu & Cập nhật Sheet thành công, màn hình sẽ hiển thị hiệu ứng pháo hoa giấy chúc mừng. Lúc này, badge trạng thái sẽ chuyển thành nút Xem Sheet màu xanh lá.")
    add_bullet_point("Hệ thống sẽ tự động tìm đúng dòng của học viên đó dựa trên Họ tên + Ngày học và cập nhật ghi đè nhận xét mới vào cột G (Nhận xét AI) mà không làm tăng thêm dòng mới.", "Trường hợp cập nhật bài viết cũ (Ghi đè): ")
    add_bullet_point("Hệ thống sẽ tự động chèn thêm một dòng dữ liệu mới vào hàng cuối cùng của bảng tính.", "Trường hợp chấm bài viết mới: ")
    add_body_text("Chị chỉ cần nhấp trực tiếp vào nút Xem Sheet này để mở nhanh file Google Sheets trên trình duyệt và kiểm soát dữ liệu.")

    # ----------------------------------------------------
    add_heading_1("IV. Giải quyết các tình huống lỗi thường gặp")
    
    add_heading_2("1. Lỗi mạng hoặc Google Sheets bị nghẽn (Hàng đợi lỗi)")
    add_body_text("Nếu trong quá trình lưu bài chấm gặp sự cố mất mạng internet hoặc link Google Sheet bị lỗi, chị đừng lo lắng. Dữ liệu bài chấm sẽ không bị mất:")
    add_bullet_point("Bài chấm bị lỗi sẽ lập tức được lưu vào ô Hàng đợi đồng bộ lỗi ở dưới cùng bên phải màn hình kèm mô tả lỗi chi tiết.", "Lưu tạm cục bộ: ")
    add_bullet_point("Sau khi mạng ổn định trở lại, chị chỉ cần bấm nút Đồng bộ lại (hình mũi tên xoay vòng) bên cạnh tên học viên đó để đẩy lại dữ liệu lên Sheet.", "Cách đồng bộ lại: ")
    
    add_heading_2("2. Muốn tùy chỉnh câu hỏi của 28 Ngày học hoặc viết lại prompt của các kỹ năng AI")
    add_body_text("Mọi nội dung đề bài và prompt hướng dẫn chấm của từng kỹ năng AI đều có thể tùy biến hoàn toàn theo phong cách giảng dạy của chị:")
    add_body_text("1. Nhấp nút Cấu hình (⚙️) ở góc trên bên phải.")
    add_body_text("2. Chuyển sang tab Kỹ năng AI để sửa văn bản prompt hướng dẫn của từng kỹ năng, hoặc chuyển sang tab Đề bài 28 Ngày để sửa nội dung các câu hỏi định hướng.")
    add_body_text("3. Bấm Lưu cấu hình để lưu lại các chỉnh sửa của chị.")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_footer = p_footer.add_run("Chúc chị Hạnh có những trải nghiệm dạy học tuyệt vời!")
    set_font(run_footer, size=10, italic=True, color=RGBColor(127, 140, 141))

    output_path = r"d:\app chấm điểm\Huong_Dan_Su_Dung_App_Cham_Diem.docx"
    doc.save(output_path)
    print("Success: Generated docx user guide file.")

if __name__ == "__main__":
    create_guide_docx()
